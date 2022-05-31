import json
import pandas as pd
import docx
from docx2pdf import convert
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

js = r"anscombe.json"
df = pd.read_json(js)
df=df.reset_index()
df = df.drop(columns = ['index'],axis = 1)
columns= df.columns.str.strip()

doc = Document() 
title = 'Название документа'
tt = doc.add_heading(title, 2)
tt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table = doc.add_table(1, df.shape[1])
table.style = 'Table Grid'
head_cells = table.rows[0].cells

for i, item in enumerate(columns):
    p = head_cells[i].paragraphs[0]
    p.add_run(item).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#создали таблицу 
table2=doc.add_table(rows=df.shape[0]+1,cols=df.shape[1])
table2.style = 'Table Grid'

for col in range(df.shape[1]):
    cell = table2.rows[col].cells
    for row in range(df.shape[0]):
         table2.cell(row,col).text=str(df.iloc[row,col]) 

path = "output_file_path.docx"
doc.save(path)
convert(path)




