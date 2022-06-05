import argparse
import pandas as pd
import ast
import docx
from docx2pdf import convert
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import json

parser = argparse.ArgumentParser(description='Json для преобразования в pdf')
parser.add_argument('json_string', type=str, help='json в виде строки')
args = parser.parse_args()

# js = args.json_string
# print(js)

# print(args.json_string)
# print(type(args.json_string))

#json_string = ast.literal_eval(args.json_string)
js1= ast.literal_eval(json.dumps(args.json_string))
# js1 = json.loads(args.json_string)
# js1 = ast.literal_eval(js1)
print(js1)
print(type(js1))

df = pd.read_json(js1)
df=df.reset_index()
df = df.drop(columns = ['index'],axis = 1) 
columns= df.columns.str.strip()

doc = Document() 
title = 'Название документа'
tt = doc.add_heading(title, 2)
tt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table = doc.add_table(1, df.shape[1])
table.style = 'Table Grid'
head_cells = table.rows[0].cells

for i, item in enumerate(columns):
    p = head_cells[i].paragraphs[0]
    p.add_run(item).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#создали таблицу 
table2=doc.add_table(rows=df.shape[0]+1,cols=df.shape[1])
table2.style = 'Table Grid'

for col in range(df.shape[1]):
    cell = table2.rows[col].cells
    for row in range(df.shape[0]):
         table2.cell(row,col).text=str(df.iloc[row,col]) 

path = "output_file_path1.docx"
doc.save(path)
convert(path)




